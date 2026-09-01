# -*- coding: utf-8 -*-
"""根据图片内容生成《作文评价量化表》Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 数据 ----------
# (一级指标, 二级指标, 一级, 二级, 三级, 四级, 对应概念)
ROWS = [
    ("条件等级", "选材", "选材虚构且普通(D)", "选材真实但事例较普通(C)", "选材真实、较新颖(B)", "选材真实、新颖(A)", ""),
    ("条件等级", "写他人", "纯属虚构(D)", "偶遇式人物(C)", "接触一段时间的人物(B)", "接触了较长时间的人物(A)", ""),
    ("条件等级", "写自己", "虚构,无依据(D分)", "所闻,有感而发(C分)", "所见,有感而发(B分)", "亲身经历(A分)", ""),
    ("过程等级", "视角", "描述的事情流水账式(D)", "描述的事情视角单一(C)", "描述事情具有视角变化(B)", "描述事情视角变化,运用合理(A)", "视角"),
    ("过程等级", "典型事例", "没有典型事例(D)", "涉及1件典型事例(C)", "涉及2件典型事例,但没有体现人物精神(B)", "涉及3件典型事例,且很好地体现了人物精神(A)", ""),
    ("过程等级", "中心", "没有突出中心(D)", "部分事件突出中心(C)", "事件突出中心(B)", "中心立意深(A)", "中心"),
    ("过程等级", "逻辑", "事件之间关联性不强,逻辑混乱(D)", "事件之间有关联,但是没有逻辑(C)", "事件之间有关联,但是逻辑不强(B)", "事件之间关联性强,逻辑清晰(A)", ""),
    ("过程等级", "详略", "无详略(D)", "详略不当(C)", "详略适当(B)", "详略得当(A)", ""),
    ("写作质量", "情节", "情节平铺直叙,平淡无奇(D)", "有矛盾点、小悬念或出人意料的叙事技巧,但运用不恰当(C)", "有矛盾点、小悬念或出人意料的叙事技巧,运用恰当(B)", "有矛盾点、小悬念或出人意料的叙事技巧,运用恰当且有一定艺术效果(A)", ""),
    ("写作质量", "写作顺序", "没有写作顺序(D)", "有一定写作顺序但稍混乱(C)", "有明确写作顺序但详略不当(B)", "有明确的写作顺序且详略得当(A)", ""),
    ("写作质量", "情感变化", "无记录情感(D)", "有记录情感(C)", "有记录两种情感变化(B)", "有记录三种或三种以上情感变化(A)", "情感变化"),
    ("写作质量", "情感描写", "无有效体现情感的描写(D)", "有两种体现情感的描写(C)", "有三种体现情感的描写(包含环境描写)(B)", "有五种或五种以上体现情感的描写(包含环境描写)(A)", "体现情感的描写"),
    ("写作质量", "抒情", "没有运用抒情(D)", "只运用了1处直接或间接抒情(C)", "运用2处抒情方式,抒发情感适中(B)", "运用3处以上抒情方式,情感真挚(A)", "抒情"),
    ("写作质量", "人物描写", "没有运用人物描写(D)", "运用了1种人物描写(C)", "运用2种人物描写(B)", "运用3种人物描写,能突出人物形象,显得真实、生动(A)", "人物描写"),
    ("写作质量", "环境描写", "没有运用环境描写(D)", "有1处环境描写(C)", "有2处环境描写(B)", "有2处环境描写且能很好地服务主题(A)", "环境描写"),
    ("写作质量", "写作手法", "没有运用写作手法(D)", "运用了1种写作手法,但并没有有效突出人物精神(C)", "运用2种写作手法,且人物精神突出(B)", "运用3种写作手法,并融入抒情与议论,人物精神清晰、突出(A)", "写作手法"),
    ("写作质量", "语言积累", "没有任何成语和叠词(D)", "有少量成语和叠词(C)", "4-5个成语和叠词并准确使用(B)", "五个以上成语、叠词并运用恰当(A)", ""),
    ("书写", "书写", "书写潦草(D)", "书写一般,有少数修改(C)", "书写工整,没有修改(B)", "书写完美,没有修改(A)", ""),
]

NOTE = "说明:一共有20项指标,一个A级相当于8分,一个B相当于5分,一个C相当于3分,一个D级相当于2分。"

# ---------- 工具函数 ----------
def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_text(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                  font_ea='宋体', font_ascii='Times New Roman', color=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_ascii
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_ea)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

# ---------- 文档 ----------
doc = Document()

# 页面: A4 横向, 窄边距
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
sec.left_margin = sec.right_margin = Cm(0.9)
sec.top_margin = sec.bottom_margin = Cm(1.0)

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("作文评价量化表")
trun.font.size = Pt(16)
trun.font.bold = True
trun.font.name = 'Times New Roman'
trun._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title.paragraph_format.space_after = Pt(2)
title.paragraph_format.space_before = Pt(0)

# 信息表头: 班级/姓名/小组/日期 填写栏
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.paragraph_format.space_after = Pt(3)
irun = info.add_run("班级:__________    姓名:__________    小组:__________    日期:__________")
irun.font.size = Pt(10.5)
irun.font.bold = True
irun.font.name = 'Times New Roman'
irun._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 表格
COL_W = [Cm(1.5), Cm(1.7), Cm(4.0), Cm(4.1), Cm(4.4), Cm(4.9), Cm(2.4), Cm(2.4), Cm(2.3)]
HEADERS = ["评价指标", "", "一", "二", "三", "四", "自评\n(评语/等级)", "小组评\n(评语/等级)", "对应概念"]

table = doc.add_table(rows=1 + len(ROWS), cols=9)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

# 表头行设为跨页重复
trPr = table.rows[0]._tr.get_or_add_trPr()
tblHeader = OxmlElement('w:tblHeader')
tblHeader.set(qn('w:val'), 'true')
trPr.append(tblHeader)

# 表头
hdr = table.rows[0]
for j, text in enumerate(HEADERS):
    c = hdr.cells[j]
    c.width = COL_W[j]
    set_cell_shading(c, "4472C4")  # 深蓝底
    for line_i, line in enumerate(text.split("\n")):
        if line_i == 0:
            set_cell_text(c, line, size=9.5, bold=True, font_ea='黑体', color='FFFFFF')
        else:
            p = c.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line)
            r.font.size = Pt(8)
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            r.font.color.rgb = RGBColor.from_string('FFFFFF')

# 合并表头"评价指标"两列
hdr.cells[0].merge(hdr.cells[1])

# 数据行
GROUP_COLORS = {"条件等级": "E2EFDA", "过程等级": "FFF2CC", "写作质量": "DEEBF7", "书写": "FCE4EC"}
for i, (g, sub, a, b, c_, d, concept) in enumerate(ROWS, start=1):
    row = table.rows[i]
    cells = row.cells
    for j, w in enumerate(COL_W):
        cells[j].width = w
    # 一级指标(后面统一合并)
    set_cell_text(cells[0], g, size=9.5, bold=True, font_ea='黑体')
    set_cell_shading(cells[0], GROUP_COLORS[g])
    # 二级指标
    set_cell_text(cells[1], sub, size=8.5, bold=True)
    set_cell_shading(cells[1], "F2F2F2")
    # 四个等级: 左对齐便于阅读
    for j, txt in zip((2, 3, 4, 5), (a, b, c_, d)):
        set_cell_text(cells[j], txt, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
    # 自评/小组评留空
    # 对应概念
    if concept:
        set_cell_text(cells[8], concept, size=8.5, bold=True)
        set_cell_shading(cells[8], "F2F2F2")

# 合并一级指标的纵向单元格
groups = {}
for i, (g, *_rest) in enumerate(ROWS, start=1):
    groups.setdefault(g, []).append(i)
for g, idxs in groups.items():
    first = table.rows[idxs[0]].cells[0]
    for k in idxs[1:]:
        first = first.merge(table.rows[k].cells[0])

# 说明
note = doc.add_paragraph()
nrun = note.add_run(NOTE)
nrun.font.size = Pt(9)
nrun.font.bold = True
nrun.font.name = 'Times New Roman'
nrun._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
note.paragraph_format.space_before = Pt(3)

out = r"C:\Users\57775\Desktop\作文评价量化表.docx"
doc.save(out)
print("saved:", out)

# TODO: 支持自定义表格配色
# 实验分支的改动
